# -*- coding: utf-8 -*-

import openpyxl
import os
import PyPDF2
import docx


os.chdir('/home/mendoza/Documentos/udemy/python/read_excel')


def debug(**Kwargs):
  for index, value in Kwargs.items():
    print('%s -> %s' %(index, value))


def debug_type(**Kwargs):
  for index, value in Kwargs.items():
    type('%s -> %s' %(index, value))


"""
  Reading Excel Documents
"""
def excel_read():
  workbook = openpyxl.load_workbook('excel_document.xlsx')
  sheet = workbook.get_sheet_by_name('Sheet1')
  # all_sheets = workbook.get_sheet_names()

  # cell = sheet['A1']
  print(str(sheet['A1'].value))

  debug(A1=str(sheet['A1'].value), B1=str(sheet['B1'].value),C1=str(sheet['C1'].value))

  # Loop
  for i in range(1, 8):
    print(i, sheet.cell(row=i, column=2).value)


"""
  Editing Excel Document
"""
def excel_edit():
  wb = openpyxl.Workbook()
  sheet1 = wb.get_sheet_by_name('Sheet')
  sheet1['A1'] = 24
  sheet1['B1'] = 'Hello'
  wb.save('excel_document2.xlsx')
  sheet2 = wb.create_sheet()
  sheet2.title = 'My new sheet name'
  wb.create_sheet(index=0, title='My first sheet renamed')
  debug(st=wb.get_sheet_names())
  wb.save('excel_document2.xlsx')


"""
  Read PDF File
  Download: http://autbor.com/meetingminutes1.pdf, http://autbor.com/meetingminutes2.pdf
"""
def pdf_read():
  pdf_file = open('meetingminutes1.pdf', 'rb')
  reader = PyPDF2.PdfFileReader(pdf_file)
  # reader.numPages
  # page = reader.getPage(0)
  # debug(pages=reader.numPages, text=page.extractText())

  for page_num in range(reader.numPages):
    debug(pages=reader.getPage(page_num).extractText())

  pdf_file.close()


def combinated_file():
  pdf1File = open('meetingminutes1.pdf', 'rb')
  pdf2File = open('meetingminutes2.pdf', 'rb')
  reader1 = PyPDF2.PdfFileReader(pdf1File)
  reader2 = PyPDF2.PdfFileReader(pdf2File)
  writer = PyPDF2.PdfFileWriter()

  for pageNum in range(reader1.numPages):
    page = reader1.getPage(pageNum)
    writer.addPage(page)

  for pageNum in range(reader2.numPages):
    page = reader2.getPage(pageNum)
    writer.addPage(page)

  outputFile = open('combinatedminutes.pdf', 'wb')
  writer.write(outputFile)
  pdf1File.close()
  pdf2File.close()


"""
  Read Word Document
"""
def word_read():
  d = docx.Document('word_document.docx')  # new instance
  p = d.paragraphs[1]
  # debug(first=d.paragraphs[0].text, second=d.paragraphs[1].text)
  # debug(a=p.runs[0].text, b=p.runs[1].text, c=p.runs[2].text, d=p.runs[3].text, f=p.runs[1].bold, g=p.runs[3].italic)

  # p.runs[3].underline = True
  # p.runs[3].text = 'Italico e sobrelinha'
  # d.save('word-document2.docx')

  # print(p.style)  # identificar estilos dos caracteres
  # d.add_paragraph('some text here')  # new paragraph
  # p.add_run('some text here') # new text on paragraph


def fetchWordText(filename):
  doc = docx.Document(filename)
  fulltext = []

  for p in doc.paragraphs:
    fulltext.append(p.text)
  
  return '\n'.join(fulltext)


debug(p=fetchWordText('word_document.docx'))